# index_builder.py
import os
from pathlib import Path
from langchain.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import SentenceTransformerEmbeddings
from langchain_community.vectorstores import Chroma

from pypdf import PdfReader
import docx
from tqdm import tqdm

DATA_DIR = Path("data")
PERSIST_DIR = "chroma_db"
CHUNK_SIZE = 800
CHUNK_OVERLAP = 200

def pdf_to_text(path: Path) -> str:
    reader = PdfReader(str(path))
    pages = []
    for p in reader.pages:
        try:
            text = p.extract_text() or ""
        except Exception:
            text = ""
        pages.append(text)
    return "\n\n".join(pages)

def docx_to_text(path: Path) -> str:
    doc = docx.Document(str(path))
    return "\n\n".join([p.text for p in doc.paragraphs])

def load_documents() -> list:
    docs = []
    for p in DATA_DIR.glob("*"):
        if p.is_dir():
            continue
        suffix = p.suffix.lower()
        try:
            if suffix == ".pdf":
                text = pdf_to_text(p)
            elif suffix in [".docx", ".doc"]:
                text = docx_to_text(p)
            elif suffix == ".txt":
                text = p.read_text(encoding="utf-8")
            else:
                print("Skipping (unknown file type):", p)
                continue
            if not text.strip():
                continue
            docs.append(Document(page_content=text, metadata={"source": p.name}))
            print("Loaded:", p.name)
        except Exception as e:
            print("Failed to read", p, e)
    return docs

def main():
    print("Looking for files in:", DATA_DIR.resolve())
    if not DATA_DIR.exists():
        print("ERROR: data/ folder not found.")
        return

    documents = load_documents()
    if not documents:
        print("No documents loaded. Put your handbook in the data/ folder.")
        return

    print(f"Loaded {len(documents)} documents. Chunking...")
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
    chunks = text_splitter.split_documents(documents)
    print(f"Split into {len(chunks)} chunks.")

    embeddings = SentenceTransformerEmbeddings(model_name="all-MiniLM-L6-v2")

    print("Building Chroma vector store (persisting to: {})".format(PERSIST_DIR))
    vectordb = Chroma.from_documents(chunks, embeddings, persist_directory=PERSIST_DIR)
    vectordb.persist()
    print("Index built and saved to:", PERSIST_DIR)

if __name__ == "__main__":
    main()
